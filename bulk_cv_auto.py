import os
import re
import random
import pandas as pd
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from openpyxl import load_workbook
from docx import Document
# Remove docx2pdf because may not work on server
from PyPDF2 import PdfMerger

# -------------------- Paths --------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

MAIN_PATH = os.path.join(BASE_DIR, "data", "MainData.xlsx")
EXP_AUTO_PATH = os.path.join(BASE_DIR, "data", "ExpAuto.xlsx")
EXP_MANUAL_PATH = os.path.join(BASE_DIR, "data", "ExpManual.xlsx")
UPDATED_PATH = os.path.join(BASE_DIR, "data", "MainData_Updated.xlsx")

EXP1_FOLDER = os.path.join(BASE_DIR, "experience", "Exp1")
EXP2_FOLDER = os.path.join(BASE_DIR, "experience", "Exp2")
CV_FOLDER   = os.path.join(BASE_DIR, "cv_templates")
TEMP_FOLDER = os.path.join(BASE_DIR, "temp")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "data", "output")

CCC_TEMPLATE = os.path.join(BASE_DIR, "data", "ccc_template.docx")

os.makedirs(TEMP_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# -------------------- Fallback PDF conversion function --------------------
import subprocess
import shutil
def docx_to_pdf_with_libreoffice(docx_path, out_dir):
    """
    Convert docx -> pdf using libreoffice CLI. Returns pdf path or None.
    """
    if not os.path.exists(docx_path):
        return None
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("[WARN] LibreOffice not found on PATH — skipping PDF conversion.")
        return None
    try:
        os.makedirs(out_dir, exist_ok=True)
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if os.path.exists(pdf_path):
            return pdf_path
        return None
    except Exception as e:
        print(f"[ERROR] LibreOffice conversion failed: {e}")
        return None

# ================================
# … (Your existing adjust_dates() and other logic unchanged) …
# =================================
def adjust_dates():
    print("[Step 1] Adjusting dates and experience info...")

    # Load experience data
    try:
        exp_auto_samples = pd.read_excel(EXP_AUTO_PATH)
    except Exception as e:
        print(f"[ERROR] Could not read ExpAuto.xlsx: {e}")
        exp_auto_samples = pd.DataFrame()

    try:
        exp_manual_samples = pd.read_excel(EXP_MANUAL_PATH)
    except Exception as e:
        print(f"[ERROR] Could not read ExpManual.xlsx: {e}")
        exp_manual_samples = pd.DataFrame()

    min_exp_years = 2
    max_exp_years = 3
    min_gap_months = 12
    max_gap_months = 36
    date_format = "%d-%m-%Y"

    df_main = pd.read_excel(MAIN_PATH)
    df_main = df_main.replace("nan", pd.NA).astype("object")
    today = datetime.today()

    def random_date_between(start, end):
        delta_days = (end - start).days
        return start if delta_days <= 0 else start + timedelta(days=random.randint(0, delta_days))

    def realistic_experience(start_date, min_years=2, max_years=3):
        years = random.randint(min_years, max_years)
        months = random.randint(0, 11)
        days = random.randint(0, 27)
        end_date = start_date + relativedelta(years=years, months=months, days=days)
        if end_date > today - timedelta(days=15):
            end_date = today - timedelta(days=random.randint(15, 90))
        return end_date

    for i in range(len(df_main)):
        manual_mode_detected = False

        if all([
            "From" in df_main.columns and pd.notna(df_main.loc[i, "From"]),
            "To" in df_main.columns and pd.notna(df_main.loc[i, "To"]),
            "From2" in df_main.columns and pd.notna(df_main.loc[i, "From2"]),
            "To2" in df_main.columns and pd.notna(df_main.loc[i, "To2"]),
            "Exp1 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp1 Company"]),
            "Exp2 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp2 Company"])
        ]):
            manual_mode_detected = True

        exp_samples = exp_manual_samples if manual_mode_detected else exp_auto_samples

        if manual_mode_detected:
            for exp_idx in [1, 2]:
                comp_col = f"Exp{exp_idx} Company"
                proj_col = f"Exp{exp_idx} Project"
                file_col = f"exp{exp_idx}"
                selected_company = str(df_main.loc[i, comp_col]).strip()
                if selected_company.lower() != "nan":
                    match = exp_samples[exp_samples["Company"].astype(str).str.strip() == selected_company]
                    if not match.empty:
                        df_main.loc[i, proj_col] = match.iloc[0]["Project"]
                        df_main.loc[i, file_col] = match.iloc[0]["File Name"]
            print(f"[INFO] Row {i+1}: Manual mode used")
            continue

        dob_raw = df_main.loc[i, "dob"] if "dob" in df_main.columns else ""
        try:
            dob = datetime.strptime(str(dob_raw).strip(), "%d-%m-%Y")
        except:
            print(f"[WARN] Row {i+1}: invalid DOB '{dob_raw}' — skipped")
            continue

        career_start = dob + relativedelta(years=18)
        if career_start.year < 2015:
            career_start = datetime(2015, 1, 1)

        total_years = (today - career_start).days / 365.0
        min_total_needed = min_exp_years*2 + min_gap_months/12.0

        if total_years < min_total_needed:
            total_days = 3*365
            gap_days = random.randint(60, 90)
            max_exp2_days = total_days - gap_days - 365
            min_exp2_days = 365
            if max_exp2_days < min_exp2_days:
                exp2_days = total_days // 2
                exp1_days = total_days - exp2_days - gap_days
            else:
                exp2_days = random.randint(min_exp2_days, max_exp2_days)
                exp1_days = total_days - gap_days - exp2_days

            exp2_end = today - timedelta(days=random.randint(15, 90))
            exp2_start = exp2_end - timedelta(days=exp2_days)
            exp1_end = exp2_start - timedelta(days=gap_days)
            exp1_start = exp1_end - timedelta(days=exp1_days)
        else:
            exp1_start = random_date_between(career_start, career_start + relativedelta(days=365))
            exp1_end = realistic_experience(exp1_start, min_exp_years, max_exp_years)
            gap_months = random.randint(min_gap_months, max_gap_months)
            exp2_start = exp1_end + relativedelta(months=gap_months)
            if exp2_start > today - timedelta(days=15):
                exp2_start = exp1_end + relativedelta(days=random.randint(30, 90))
            exp2_end = realistic_experience(exp2_start, min_exp_years, max_exp_years)
            if exp2_start > exp2_end:
                exp2_start = exp2_end - relativedelta(years=min_exp_years)

        try:
            samples = exp_samples.sample(2).reset_index(drop=True)
            df_main.loc[i, "Exp1 Company"] = str(samples.loc[0, "Company"])
            df_main.loc[i, "Exp1 Project"] = str(samples.loc[0, "Project"])
            df_main.loc[i, "From"] = exp1_start.strftime(date_format)
            df_main.loc[i, "To"] = exp1_end.strftime(date_format)
            df_main.loc[i, "Exp2 Company"] = str(samples.loc[1, "Company"])
            df_main.loc[i, "Exp2 Project"] = str(samples.loc[1, "Project"])
            df_main.loc[i, "From2"] = exp2_start.strftime(date_format)
            df_main.loc[i, "To2"] = exp2_end.strftime(date_format)
            if "File Name" in samples.columns:
                df_main.loc[i, "exp1"] = str(samples.loc[0, "File Name"])
                df_main.loc[i, "exp2"] = str(samples.loc[1, "File Name"])
        except Exception as ex:
            print(f"[WARN] Row {i+1}: Could not sample experience — {ex}")

    df_main.to_excel(UPDATED_PATH, index=False)
    print(f"[OK] Updated Excel saved at: {UPDATED_PATH}")

# ==========================================================
# STEP 2: Fill DOCX + CCC + Merge
# ==========================================================
INVALID_FILENAME_CHARS = r'[:<>"/\\|?*\n\r\t]'
def sanitize_filename(name: str) -> str:
    name = str(name).strip()
    name = re.sub(INVALID_FILENAME_CHARS, "_", name)
    return re.sub(r"\s+", " ", name)

def replace_in_paragraph(paragraph, replacements):
    full_text = "".join([r.text for r in paragraph.runs])
    new_text = full_text
    for key, val in replacements.items():
        placeholder = f"<{key}>"
        replacement_text = "" if pd.isna(val) else str(val)
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, replacement_text)
    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

def replace_placeholders(doc_path, replacements, output_path):
    try:
        doc = Document(doc_path)
        for para in doc.paragraphs:
            replace_in_paragraph(para, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, replacements)
        doc.save(output_path)
    except Exception as e:
        print(f"Error processing {doc_path}: {e}")

def merge_pdfs(pdf_list, output_path):
    merger = PdfMerger()
    for pdf in pdf_list:
        if os.path.exists(pdf):
            merger.append(pdf)
    merger.write(output_path)
    merger.close()

def get_unique_filename(base_path):
    if not os.path.exists(base_path):
        return base_path
    base, ext = os.path.splitext(base_path)
    counter = 1
    while True:
        new_path = f"{base}-{counter}{ext}"
        if not os.path.exists(new_path):
            return new_path
        counter += 1

# --- NEW CCC FORM FILLING ---
def fill_ccc_docx(data_dict, output_path):
    try:
        if not os.path.exists(CCC_TEMPLATE):
            print(f"[WARN] CCC template not found: {CCC_TEMPLATE}")
            return None

        # --- Calculate total experience ---
        try:
            fmt = "%d-%m-%Y"
            f1 = datetime.strptime(str(data_dict.get("From", "")), fmt)
            t1 = datetime.strptime(str(data_dict.get("To", "")), fmt)
            f2 = datetime.strptime(str(data_dict.get("From2", "")), fmt)
            t2 = datetime.strptime(str(data_dict.get("To2", "")), fmt)
            total_months = ((t1 - f1).days + (t2 - f2).days) // 30
            years, months = divmod(total_months, 12)
            total_exp_str = f"{years}Y{months}M"
        except Exception:
            total_exp_str = ""

        replacements = {col: val for col, val in data_dict.items() if pd.notna(val)}
        replacements["total"] = total_exp_str  # for <total>

        # --- Fill placeholders in CCC template ---
        out_docx = output_path.replace(".pdf", ".docx")
        replace_placeholders(CCC_TEMPLATE, replacements, out_docx)

        # --- Convert to PDF ---
        from docx2pdf import convert
        convert(out_docx)
        pdf_path = out_docx.replace(".docx", ".pdf")
        print(f"[OK] CCC DOCX filled and converted: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"[ERROR] Creating CCC DOCX failed: {e}")
        return None

# In your make_pdfs() function, change the conversion section to something like:

# ==========================================================
# STEP 2: Fill templates, convert to PDF, and merge
# ==========================================================
def make_pdfs():
    print("[Step 2] Filling templates, converting to PDF, and adding CCC if needed…")

    # Helper function to find file ignoring case
    def find_file_case_insensitive(folder, filename):
        """Find a file in the folder ignoring case differences."""
        if not os.path.exists(folder):
            return None
        target = filename.lower()
        for f in os.listdir(folder):
            if f.lower() == target:
                return os.path.join(folder, f)
        return None

    try:
        df = pd.read_excel(UPDATED_EXCEL_PATH)
    except Exception as e:
        print(f"[ERROR] Could not read Excel: {e}")
        return

    for idx, row in df.iterrows():
        name = str(row.get("name", "Unknown"))
        print(f"[{idx + 1}/{len(df)}] Preparing docs for: {name}")

        # Find files ignoring case
        exp1_file = find_file_case_insensitive(EXP1_FOLDER, str(row.get("exp1", "")).strip() + ".docx")
        exp2_file = find_file_case_insensitive(EXP2_FOLDER, str(row.get("exp2", "")).strip() + ".docx")
        cv_file   = find_file_case_insensitive(CV_FOLDER,   str(row.get("cv", "")).strip() + ".docx")

        # Create output paths
        temp_dir = os.path.join("temp", name)
        os.makedirs(temp_dir, exist_ok=True)
        final_docx = os.path.join(temp_dir, f"{name}_cv.docx")

        replacements = {col: str(val) for col, val in row.items() if pd.notna(val)}

        for src in [exp1_file, exp2_file, cv_file]:
            if src and os.path.exists(src):
                dest = os.path.join(temp_dir, os.path.basename(src))
                replace_placeholders(src, replacements, dest)
            else:
                print(f"[WARN] Missing source file: {src or 'Unknown'}")

        # Check if LibreOffice exists for DOCX→PDF conversion
        libreoffice_path = shutil.which("soffice")
        if not libreoffice_path:
            print("[WARN] LibreOffice not found on PATH — skipping PDF conversion.")
            continue

        # Convert DOCX → PDF
        try:
            subprocess.run([
                libreoffice_path, "--headless", "--convert-to", "pdf",
                "--outdir", temp_dir, final_docx
            ], check=True)
        except Exception as e:
            print(f"[WARN] Did not convert to PDF: {final_docx} ({e})")

        # Merge PDFs if available
        pdf_files = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir) if f.endswith(".pdf")]
        if not pdf_files:
            print(f"[WARN] No PDFs to merge for: {name}")
            continue

        output_pdf = os.path.join("output", f"{name}.pdf")
        os.makedirs("output", exist_ok=True)
        merge_pdfs(pdf_files, output_pdf)
        print(f"[OK] PDF created: {output_pdf}")

        # Clean up temporary directory
        shutil.rmtree(temp_dir, ignore_errors=True)

    print("[DONE] All documents processed and temp cleaned up.")


if __name__ == "__main__":
    adjust_dates()
    make_pdfs()
