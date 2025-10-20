import os
import re
import random
import pandas as pd
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from openpyxl import load_workbook
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfMerger

# -------------------- Paths --------------------
MAIN_PATH = r"F:\CV_SOFTWARE\MainData.xlsx"
EXP_AUTO_PATH = r"F:\CV_SOFTWARE\ExpAuto.xlsx"
EXP_MANUAL_PATH = r"F:\CV_SOFTWARE\ExpManual.xlsx"
UPDATED_PATH = r"F:\CV_SOFTWARE\MainData_Updated.xlsx"

EXP1_FOLDER = r"F:\CV_SOFTWARE\Experience Letters\Exp1"
EXP2_FOLDER = r"F:\CV_SOFTWARE\Experience Letters\Exp2"
CV_FOLDER   = r"F:\CV_SOFTWARE\CVs"
TEMP_FOLDER = r"F:\CV_SOFTWARE\temp"
OUTPUT_FOLDER = r"F:\CV_SOFTWARE\output"

os.makedirs(TEMP_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# -------------------- Load Experience Samples --------------------
exp_auto_samples = pd.read_excel(EXP_AUTO_PATH) if os.path.exists(EXP_AUTO_PATH) else None
exp_manual_samples = pd.read_excel(EXP_MANUAL_PATH) if os.path.exists(EXP_MANUAL_PATH) else None
if exp_auto_samples is None:
    raise FileNotFoundError(f"Auto experience file not found: {EXP_AUTO_PATH}")
if exp_manual_samples is None:
    raise FileNotFoundError(f"Manual experience file not found: {EXP_MANUAL_PATH}")

# -------------------- STEP 1: Adjust Dates --------------------
def adjust_dates():
    print("[Step 1] Adjusting dates and experience info...")

    min_exp_years = 2
    max_exp_years = 3
    min_gap_months = 12
    max_gap_months = 36
    date_format = "%d-%m-%Y"

    df_main = pd.read_excel(MAIN_PATH)
    df_main = df_main.replace("nan", pd.NA).astype("object")
    today = datetime.today()

    def random_date_between(start, end):
        delta_days = (end - start).days
        return start if delta_days <= 0 else start + timedelta(days=random.randint(0, delta_days))

    def realistic_experience(start_date, min_years=2, max_years=3):
        years = random.randint(min_years, max_years)
        months = random.randint(0, 11)
        days = random.randint(0, 27)
        end_date = start_date + relativedelta(years=years, months=months, days=days)
        if end_date > today - timedelta(days=15):
            end_date = today - timedelta(days=random.randint(15, 90))
        return end_date

    for i in range(len(df_main)):
        manual_mode_detected = False

        # -------------------- Detect Manual Mode --------------------
        if all([
            "From" in df_main.columns and pd.notna(df_main.loc[i, "From"]),
            "To" in df_main.columns and pd.notna(df_main.loc[i, "To"]),
            "From2" in df_main.columns and pd.notna(df_main.loc[i, "From2"]),
            "To2" in df_main.columns and pd.notna(df_main.loc[i, "To2"]),
            "Exp1 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp1 Company"]),
            "Exp2 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp2 Company"])
        ]):
            manual_mode_detected = True

        # -------------------- Select Samples --------------------
        exp_samples = exp_manual_samples if manual_mode_detected else exp_auto_samples

        # -------------------- Manual Mode --------------------
        if manual_mode_detected:
            for exp_idx in [1, 2]:
                comp_col = f"Exp{exp_idx} Company"
                proj_col = f"Exp{exp_idx} Project"
                file_col = f"exp{exp_idx}"
                selected_company = str(df_main.loc[i, comp_col]).strip()
                if selected_company.lower() != "nan":
                    match = exp_samples[exp_samples["Company"].astype(str).str.strip() == selected_company]
                    if not match.empty:
                        df_main.loc[i, proj_col] = match.iloc[0]["Project"]
                        df_main.loc[i, file_col] = match.iloc[0]["File Name"]
            print(f"[INFO] Row {i+1}: Manual mode used")
            continue

        # -------------------- Auto Mode --------------------
        dob_raw = df_main.loc[i, "dob"] if "dob" in df_main.columns else ""
        try:
            dob = datetime.strptime(str(dob_raw).strip(), "%d-%m-%Y")
        except:
            print(f"[WARN] Row {i+1}: invalid DOB '{dob_raw}' — skipped")
            continue

        career_start = dob + relativedelta(years=18)
        if career_start.year < 2015:
            career_start = datetime(2015, 1, 1)

        total_years = (today - career_start).days / 365.0
        min_total_needed = min_exp_years*2 + min_gap_months/12.0

        # -------------------- Short history (< 2–3 years) --------------------
        if total_years < min_total_needed:
            total_days = 3*365
            gap_days = random.randint(60, 90)
            max_exp2_days = total_days - gap_days - 365
            min_exp2_days = 365
            if max_exp2_days < min_exp2_days:
                exp2_days = total_days // 2
                exp1_days = total_days - exp2_days - gap_days
            else:
                exp2_days = random.randint(min_exp2_days, max_exp2_days)
                exp1_days = total_days - gap_days - exp2_days

            exp2_end = today - timedelta(days=random.randint(15, 90))
            exp2_start = exp2_end - timedelta(days=exp2_days)
            exp1_end = exp2_start - timedelta(days=gap_days)
            exp1_start = exp1_end - timedelta(days=exp1_days)
        else:
            exp1_start = random_date_between(career_start, career_start + relativedelta(days=365))
            exp1_end = realistic_experience(exp1_start, min_exp_years, max_exp_years)
            gap_months = random.randint(min_gap_months, max_gap_months)
            exp2_start = exp1_end + relativedelta(months=gap_months)
            if exp2_start > today - timedelta(days=15):
                exp2_start = exp1_end + relativedelta(days=random.randint(30, 90))
            exp2_end = realistic_experience(exp2_start, min_exp_years, max_exp_years)
            if exp2_start > exp2_end:
                exp2_start = exp2_end - relativedelta(years=min_exp_years)

        # -------------------- Assign Experiences --------------------
        try:
            samples = exp_samples.sample(2).reset_index(drop=True)
            df_main.loc[i, "Exp1 Company"] = str(samples.loc[0, "Company"])
            df_main.loc[i, "Exp1 Project"] = str(samples.loc[0, "Project"])
            df_main.loc[i, "From"] = exp1_start.strftime(date_format)
            df_main.loc[i, "To"] = exp1_end.strftime(date_format)
            df_main.loc[i, "Exp2 Company"] = str(samples.loc[1, "Company"])
            df_main.loc[i, "Exp2 Project"] = str(samples.loc[1, "Project"])
            df_main.loc[i, "From2"] = exp2_start.strftime(date_format)
            df_main.loc[i, "To2"] = exp2_end.strftime(date_format)
            if "File Name" in samples.columns:
                df_main.loc[i, "exp1"] = str(samples.loc[0, "File Name"])
                df_main.loc[i, "exp2"] = str(samples.loc[1, "File Name"])
        except Exception as ex:
            print(f"[WARN] Row {i+1}: Could not sample experience — {ex}")

    df_main.to_excel(UPDATED_PATH, index=False)
    print(f"[OK] Updated Excel saved at: {UPDATED_PATH}")

    # -------------------- Fix Width --------------------
    try:
        wb = load_workbook(UPDATED_PATH)
        ws = wb.active
        for col in ws.columns:
            max_len = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = max_len + 2
        wb.save(UPDATED_PATH)
        wb.close()
    except Exception as e:
        print(f"[WARN] Column width adjust failed: {e}")

    print(f"[OK] Dates updated and saved to: {UPDATED_PATH}")

# ==========================================================
# STEP 2: Fill DOCX templates and make PDFs
# ==========================================================
INVALID_FILENAME_CHARS = r'[:<>"/\\|?*\n\r\t]'

def sanitize_filename(name: str) -> str:
    name = str(name).strip()
    name = re.sub(INVALID_FILENAME_CHARS, "_", name)
    return re.sub(r"\s+", " ", name)

# --- your original formatting-preserving functions ---
def replace_in_paragraph(paragraph, replacements):
    full_text = "".join([r.text for r in paragraph.runs])
    new_text = full_text
    for key, val in replacements.items():
        placeholder = f"<{key}>"
        replacement_text = "" if pd.isna(val) else str(val)
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, replacement_text)

    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

def replace_placeholders(doc_path, replacements, output_path):
    try:
        doc = Document(doc_path)
        for para in doc.paragraphs:
            replace_in_paragraph(para, replacements)
        for table in doc.tables:
            for row in table.rows:
                try:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_in_paragraph(para, replacements)
                except ValueError as e:
                    print(f"Skipped merged cell in '{doc_path}': {e}")
        doc.save(output_path)
    except Exception as e:
        print(f"Error processing {doc_path}: {e}")

def merge_pdfs(pdf_list, output_path):
    merger = PdfMerger()
    for pdf in pdf_list:
        if os.path.exists(pdf):
            merger.append(pdf)
    merger.write(output_path)
    merger.close()

def get_unique_filename(base_path):
    if not os.path.exists(base_path):
        return base_path
    base, ext = os.path.splitext(base_path)
    counter = 1
    while True:
        new_path = f"{base}-{counter}{ext}"
        if not os.path.exists(new_path):
            return new_path
        counter += 1

def make_pdfs():
    print("[Step 2] Filling templates and converting to PDF...")

    df = pd.read_excel(UPDATED_PATH)
    names_list = []

    for idx, row in df.iterrows():
        name = sanitize_filename(row.get("Name", f"row_{idx+1}"))
        print(f"[{idx+1}/{len(df)}] Preparing docs for: {name}")
        replacements = {col: row[col] for col in df.columns if pd.notna(row[col])}

        exp1_file = os.path.join(EXP1_FOLDER, str(row.get("exp1", "")).strip() + ".docx")
        exp2_file = os.path.join(EXP2_FOLDER, str(row.get("exp2", "")).strip() + ".docx")
        cv_file   = os.path.join(CV_FOLDER,   str(row.get("cv", "")).strip()   + ".docx")

        for label, src in [("exp1", exp1_file), ("exp2", exp2_file), ("cv", cv_file)]:
            if os.path.exists(src):
                out_doc = os.path.join(TEMP_FOLDER, f"{name}_{label}.docx")
                replace_placeholders(src, replacements, out_doc)
            else:
                print(f"[WARN] Missing {label.upper()} for {name}")

        names_list.append(name)

    print("[INFO] Bulk converting DOCX to PDF...")
    convert(TEMP_FOLDER)

    for name in names_list:
        pdfs = [
            os.path.join(TEMP_FOLDER, f"{name}_cv.pdf"),
            os.path.join(TEMP_FOLDER, f"{name}_exp1.pdf"),
            os.path.join(TEMP_FOLDER, f"{name}_exp2.pdf"),
        ]

        # Define the output folder path and desktop path separately
        final_pdf_output_folder = os.path.join(OUTPUT_FOLDER, f"{name}.pdf")
        final_pdf_desktop = os.path.join("C:\\Users\\sidetable\\Desktop", f"{name}.pdf")

        # Get unique filenames for each location
        final_pdf_output_folder = get_unique_filename(final_pdf_output_folder)
        final_pdf_desktop = get_unique_filename(final_pdf_desktop)

        # Check if PDFs exist
        existing = [p for p in pdfs if os.path.exists(p)]
        if existing:
            # Merge PDFs and save to OUTPUT_FOLDER
            merge_pdfs(existing, final_pdf_output_folder)

            # Merge PDFs and save to Desktop
            merge_pdfs(existing, final_pdf_desktop)

            print(f"[OK] Created in OUTPUT_FOLDER: {final_pdf_output_folder}")
            print(f"[OK] Created on Desktop: {final_pdf_desktop}")
        else:
            print(f"[WARN] No PDFs found for {name}")

    #Clear Temp Folder
    for f in os.listdir(TEMP_FOLDER):
        if f.lower().endswith((".docx", ".pdf")):
            try:
                os.remove(os.path.join(TEMP_FOLDER, f))
            except Exception as e:
                print(f"[WARN] Cleanup failed for {f}: {e}")

    print("[DONE] All documents processed and temp cleaned up.")

# ==========================================================
# RUN BOTH STEPS
# ==========================================================
if __name__ == "__main__":
    adjust_dates()
    make_pdfs()
