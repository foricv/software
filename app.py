import os
import re
import io
import random
import threading
import time
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta

from flask import Flask, render_template, request, redirect, jsonify, Response, send_file
import pandas as pd
from openpyxl import load_workbook
from docx import Document

# -------------------- Flask App --------------------
app = Flask(__name__)

# -------------------- Relative Paths --------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MAIN_PATH = os.path.join(BASE_DIR, "data", "MainData.xlsx")
EXP_AUTO_PATH = os.path.join(BASE_DIR, "data", "ExpAuto.xlsx")
EXP_MANUAL_PATH = os.path.join(BASE_DIR, "data", "ExpManual.xlsx")
CCC_TEMPLATE = os.path.join(BASE_DIR, "ccc_template.docx")
EXP1_FOLDER = os.path.join(BASE_DIR, "experience", "Exp1")
EXP2_FOLDER = os.path.join(BASE_DIR, "experience", "Exp2")
CV_FOLDER   = os.path.join(BASE_DIR, "cv_templates")
TEMP_FOLDER = os.path.join(BASE_DIR, "temp")
os.makedirs(TEMP_FOLDER, exist_ok=True)

logs = []
process_running = False

# -------------------- Helpers --------------------
INVALID_FILENAME_CHARS = r'[:<>"/\\|?*\n\r\t]'
def sanitize_filename(name: str) -> str:
    name = str(name).strip()
    name = re.sub(INVALID_FILENAME_CHARS, "_", name)
    return re.sub(r"\s+", " ", name)

def replace_in_paragraph(paragraph, replacements):
    full_text = "".join([r.text for r in paragraph.runs])
    new_text = full_text
    for key, val in replacements.items():
        placeholder = f"<{key}>"
        replacement_text = "" if pd.isna(val) else str(val)
        new_text = new_text.replace(placeholder, replacement_text)
    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

def replace_placeholders(doc_path, replacements):
    doc = Document(doc_path)
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)
    return doc

def merge_docx_files(docs_list):
    merged_doc = Document()
    merged_doc._body.clear_content()
    for doc in docs_list:
        for element in doc.element.body:
            merged_doc.element.body.append(element)
    return merged_doc

# -------------------- Adjust Dates --------------------
def adjust_dates(df_main):
    # Force all relevant columns to object (string-compatible)
    cols_to_fix = [
        "Exp1 Company", "Exp1 Project", "From", "To",
        "Exp2 Company", "Exp2 Project", "From2", "To2",
        "exp1", "exp2"
    ]
    for col in cols_to_fix:
        if col in df_main.columns:
            df_main[col] = df_main[col].astype("object")

    min_exp_years = 2
    max_exp_years = 3
    min_gap_months = 12
    max_gap_months = 36
    date_format = "%d-%m-%Y"
    today = datetime.today()

    exp_auto_samples = pd.read_excel(EXP_AUTO_PATH)
    exp_manual_samples = pd.read_excel(EXP_MANUAL_PATH)

    for i in range(len(df_main)):
        manual_mode_detected = False
        if all([
            "From" in df_main.columns and pd.notna(df_main.loc[i, "From"]),
            "To" in df_main.columns and pd.notna(df_main.loc[i, "To"]),
            "From2" in df_main.columns and pd.notna(df_main.loc[i, "From2"]),
            "To2" in df_main.columns and pd.notna(df_main.loc[i, "To2"]),
            "Exp1 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp1 Company"]),
            "Exp2 Company" in df_main.columns and pd.notna(df_main.loc[i, "Exp2 Company"])
        ]):
            manual_mode_detected = True

        exp_samples = exp_manual_samples if manual_mode_detected else exp_auto_samples

        if manual_mode_detected:
            for exp_idx in [1, 2]:
                comp_col = f"Exp{exp_idx} Company"
                proj_col = f"Exp{exp_idx} Project"
                file_col = f"exp{exp_idx}"
                selected_company = str(df_main.loc[i, comp_col]).strip()
                if selected_company.lower() != "nan":
                    match = exp_samples[exp_samples["Company"].astype(str).str.strip() == selected_company]
                    if not match.empty:
                        df_main.loc[i, proj_col] = match.iloc[0]["Project"]
                        df_main.loc[i, file_col] = match.iloc[0]["File Name"]
            continue

        dob_raw = df_main.loc[i, "dob"] if "dob" in df_main.columns else ""
        try:
            dob = datetime.strptime(str(dob_raw).strip(), "%d-%m-%Y")
        except:
            continue

        career_start = dob + relativedelta(years=18)
        if career_start.year < 2015:
            career_start = datetime(2015, 1, 1)

        def random_date_between(start, end):
            delta_days = (end - start).days
            return start if delta_days <= 0 else start + timedelta(days=random.randint(0, delta_days))

        def realistic_experience(start_date, min_years=2, max_years=3):
            years = random.randint(min_years, max_years)
            months = random.randint(0, 11)
            days = random.randint(0, 27)
            end_date = start_date + relativedelta(years=years, months=months, days=days)
            if end_date > today - timedelta(days=15):
                end_date = today - timedelta(days=random.randint(15, 90))
            return end_date

        exp1_start = random_date_between(career_start, career_start + relativedelta(days=365))
        exp1_end = realistic_experience(exp1_start, min_exp_years, max_exp_years)
        gap_months = random.randint(min_gap_months, max_gap_months)
        exp2_start = exp1_end + relativedelta(months=gap_months)
        if exp2_start > today - timedelta(days=15):
            exp2_start = exp1_end + relativedelta(days=random.randint(30, 90))
        exp2_end = realistic_experience(exp2_start, min_exp_years, max_exp_years)

        try:
            samples = exp_samples.sample(2).reset_index(drop=True)
            df_main.loc[i, "Exp1 Company"] = str(samples.loc[0, "Company"])
            df_main.loc[i, "Exp1 Project"] = str(samples.loc[0, "Project"])
            df_main.loc[i, "From"] = exp1_start.strftime(date_format)
            df_main.loc[i, "To"] = exp1_end.strftime(date_format)
            df_main.loc[i, "Exp2 Company"] = str(samples.loc[1, "Company"])
            df_main.loc[i, "Exp2 Project"] = str(samples.loc[1, "Project"])
            df_main.loc[i, "From2"] = exp2_start.strftime(date_format)
            df_main.loc[i, "To2"] = exp2_end.strftime(date_format)
            if "File Name" in samples.columns:
                df_main.loc[i, "exp1"] = str(samples.loc[0, "File Name"])
                df_main.loc[i, "exp2"] = str(samples.loc[1, "File Name"])
        except:
            continue

    return df_main

# -------------------- Generate Merged DOCX --------------------
def generate_merged_docx(df_main, candidate_idx=0):
    row = df_main.iloc[candidate_idx]
    name = sanitize_filename(row.get("Name", f"candidate_{candidate_idx+1}"))
    replacements = {col: row[col] for col in df_main.columns if pd.notna(row[col])}

    # Calculate total experience
    try:
        fmt = "%d-%m-%Y"
        f1 = datetime.strptime(str(row.get("From", "")), fmt)
        t1 = datetime.strptime(str(row.get("To", "")), fmt)
        f2 = datetime.strptime(str(row.get("From2", "")), fmt)
        t2 = datetime.strptime(str(row.get("To2", "")), fmt)
        total_months = ((t1 - f1).days + (t2 - f2).days) // 30
        years, months = divmod(total_months, 12)
        replacements["total"] = f"{years}Y{months}M"
    except:
        replacements["total"] = ""

    docs_to_merge = []
    for col, folder in [("cv", CV_FOLDER), ("exp1", EXP1_FOLDER), ("exp2", EXP2_FOLDER)]:
        file_name = str(row.get(col, "")).strip()
        path = os.path.join(folder, f"{file_name}.docx")
        if os.path.exists(path):
            docs_to_merge.append(replace_placeholders(path, replacements))

    ccc_value = str(row.get("ccc", "No")).strip().lower()
    if ccc_value == "yes" and os.path.exists(CCC_TEMPLATE):
        docs_to_merge.append(replace_placeholders(CCC_TEMPLATE, replacements))

    merged_doc = merge_docx_files(docs_to_merge)
    output_stream = io.BytesIO()
    merged_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream, f"{name}_merged.docx"

def generate_individual_docs(df_main, candidate_idx=0):
    row = df_main.iloc[candidate_idx]
    name = sanitize_filename(row.get("Name", f"candidate_{candidate_idx+1}"))
    replacements = {col: row[col] for col in df_main.columns if pd.notna(row[col])}

    output_docs = []

    # Generate DOCX for each section separately
    for col, folder in [("cv", CV_FOLDER), ("exp1", EXP1_FOLDER), ("exp2", EXP2_FOLDER)]:
        file_name = str(row.get(col, "")).strip()
        path = os.path.join(folder, f"{file_name}.docx")
        if os.path.exists(path):
            doc = replace_placeholders(path, replacements)
            out_stream = io.BytesIO()
            doc.save(out_stream)
            out_stream.seek(0)
            output_docs.append((f"{name}_{col}.docx", out_stream))

    ccc_value = str(row.get("ccc", "No")).strip().lower()
    if ccc_value == "yes" and os.path.exists(CCC_TEMPLATE):
        doc = replace_placeholders(CCC_TEMPLATE, replacements)
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        output_docs.append((f"{name}_ccc.docx", out_stream))

    return output_docs

# -------------------- Routes --------------------
@app.route('/')
def index():
    try:
        df = pd.read_excel(MAIN_PATH)
        records = df.to_dict(orient='records')
        columns = df.columns.tolist()
    except:
        records, columns = [], []
    return render_template('form.html', records=records, columns=columns)

@app.route('/exp-samples')
def exp_samples_route():
    try:
        df = pd.read_excel(EXP_MANUAL_PATH).fillna("")
        required_cols = ["File Name", "Country", "Company Type", "Company", "Project"]
        rows = df[required_cols].to_dict(orient="records")
        return jsonify({"status": "ok", "rows": rows})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/submit', methods=['POST'])
def submit():
    data = request.form.to_dict()
    try:
        df = pd.read_excel(MAIN_PATH)
    except FileNotFoundError:
        df = pd.DataFrame(columns=data.keys())
    df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
    df.to_excel(MAIN_PATH, index=False)
    return redirect('/')

@app.route('/clear-data', methods=['POST'])
def clear_data():
    try:
        if os.path.exists(MAIN_PATH):
            df = pd.read_excel(MAIN_PATH)
            empty_df = pd.DataFrame(columns=df.columns)
            empty_df.to_excel(MAIN_PATH, index=False)
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/generate-docx')
def generate_docx_route():
    try:
        df_main = pd.read_excel(MAIN_PATH)
        df_main = adjust_dates(df_main)
        doc_stream, filename = generate_merged_docx(df_main)
        return send_file(
            doc_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

from flask import send_file
import zipfile
import io

@app.route('/generate-zip')

def generate_zip_route():
    try:
        df_main = pd.read_excel(MAIN_PATH)
        df_main = adjust_dates(df_main)

        zip_stream = io.BytesIO()
        with zipfile.ZipFile(zip_stream, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            for idx in range(len(df_main)):
                individual_docs = generate_individual_docs(df_main, idx)
                for filename, doc_stream in individual_docs:
                    zf.writestr(filename, doc_stream.getvalue())

        zip_stream.seek(0)
        return send_file(
            zip_stream,
            as_attachment=True,
            download_name="all_candidates.zip",
            mimetype="application/zip"
        )
    except Exception as e:
        return f"❌ Error generating ZIP: {str(e)}"



# -------------------- Run App --------------------
if __name__ == '__main__':
    app.run(debug=True)
